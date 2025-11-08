"""
Main report generation service - orchestrates all components
"""
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List
import markdown
from io import BytesIO

from backend.models.report_models import (
    ReportRequest,
    ReportResponse,
    ReportMetadata,
    ReportFormat,
    TestResult,
)
from backend.services.ai_report_service import AIReportService
from backend.services.template_service import TemplateService
from backend.services.quality_service import QualityService
from backend.services.data_extraction_service import DataExtractionService
from backend.config import get_settings


class ReportGenerationService:
    """
    Main service for generating reports in various formats
    """

    def __init__(self):
        self.settings = get_settings()
        self.ai_service = AIReportService()
        self.template_service = TemplateService()
        self.quality_service = QualityService()
        self.data_service = DataExtractionService()

    async def generate_report(self, request: ReportRequest) -> ReportResponse:
        """
        Generate report based on request

        Args:
            request: Report generation request

        Returns:
            ReportResponse with generated files
        """
        start_time = time.time()
        report_id = str(uuid.uuid4())[:8]

        try:
            # Step 1: Extract additional data if needed
            if request.excel_file:
                excel_data = await self.data_service.extract_from_excel(
                    request.excel_file
                )
                # Merge with test results
                # (implementation depends on data structure)

            # Step 2: Enhance test results with AI if enabled
            if request.enable_ai_enhancement:
                await self._enhance_test_results(request.test_results)

            # Step 3: Generate executive summary
            executive_summary = await self.ai_service.generate_executive_summary(
                [result.dict() for result in request.test_results]
            )

            # Step 4: Prepare template context
            context = self._prepare_context(request, executive_summary, report_id)

            # Step 5: Render template
            template_id = request.template_id or self._get_default_template(
                request.report_type
            )
            report_content = self.template_service.render_template(
                template_id, context
            )

            # Step 6: Quality check if enabled
            quality_result = None
            if request.enable_spell_check or request.enable_grammar_check:
                quality_result = await self.quality_service.perform_quality_check(
                    report_content, request.test_results
                )

            # Step 7: Generate output files in requested formats
            output_dir = Path(self.settings.reports_output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)

            generated_files = {}
            file_sizes = {}

            for output_format in request.output_formats:
                file_path = await self._generate_output_file(
                    report_content,
                    output_format,
                    report_id,
                    request.report_title,
                    output_dir,
                    request.test_results,
                )

                if file_path:
                    generated_files[output_format.value] = str(file_path)
                    file_sizes[output_format.value] = file_path.stat().st_size

            # Step 8: Create metadata
            metadata = ReportMetadata(
                report_id=report_id,
                generation_time_seconds=time.time() - start_time,
                ai_enhanced=request.enable_ai_enhancement,
                quality_checked=request.enable_spell_check
                or request.enable_grammar_check,
                file_paths=generated_files,
                file_sizes=file_sizes,
                quality_check_results=quality_result,
            )

            # Step 9: Create response
            response = ReportResponse(
                success=True,
                report_id=report_id,
                message=f"Report generated successfully in {len(generated_files)} format(s)",
                files=generated_files,
                metadata=metadata,
                quality_check=quality_result,
            )

            return response

        except Exception as e:
            return ReportResponse(
                success=False,
                report_id=report_id,
                message=f"Report generation failed: {str(e)}",
                metadata=ReportMetadata(
                    report_id=report_id,
                    generation_time_seconds=time.time() - start_time,
                ),
                errors=[str(e)],
            )

    async def _enhance_test_results(self, test_results: List[TestResult]):
        """
        Enhance test results with AI-generated interpretations

        Args:
            test_results: List of test results to enhance
        """
        for test in test_results:
            # Generate interpretation
            interpretation = await self.ai_service.interpret_test_results(
                test.test_name, test.measurements, test.pass_fail_criteria
            )

            # Add to test result (we'll add this field dynamically)
            test.notes = f"{test.notes}\n\n{interpretation}" if test.notes else interpretation

    def _prepare_context(
        self, request: ReportRequest, executive_summary: str, report_id: str
    ) -> Dict[str, Any]:
        """
        Prepare template rendering context

        Args:
            request: Report request
            executive_summary: Generated executive summary
            report_id: Generated report ID

        Returns:
            Context dictionary
        """
        context = {
            "report_id": report_id,
            "report_date": request.report_date.strftime("%Y-%m-%d"),
            "report_version": "1.0",
            "report_title": request.report_title,
            "client_name": request.client_name,
            "client_address": request.client_address,
            "project_name": request.project_name,
            "executive_summary": executive_summary,
            "test_results": [self._format_test_result(tr) for tr in request.test_results],
            "conclusions": "Test results are documented above. Please review the individual test sections for detailed findings.",
        }

        # Add test-specific data for performance reports
        if request.test_results:
            first_test = request.test_results[0]
            context.update({
                "sample_id": first_test.sample_id,
                "manufacturer": first_test.manufacturer,
                "model": first_test.model,
                "test_date": first_test.test_date.strftime("%Y-%m-%d"),
            })

            # Add I-V parameters if available
            if "Voc" in first_test.measurements:
                context["voc"] = first_test.measurements.get("Voc", "N/A")
                context["isc"] = first_test.measurements.get("Isc", "N/A")
                context["pmax"] = first_test.measurements.get("Pmax", "N/A")
                context["vmp"] = first_test.measurements.get("Vmp", "N/A")
                context["imp"] = first_test.measurements.get("Imp", "N/A")
                context["fill_factor"] = first_test.measurements.get("FF", "N/A")

        # Add custom fields
        context.update(request.custom_fields)

        return context

    def _format_test_result(self, test_result: TestResult) -> Dict[str, Any]:
        """
        Format test result for template rendering

        Args:
            test_result: Test result object

        Returns:
            Formatted dictionary
        """
        return {
            "test_id": test_result.test_id,
            "test_name": test_result.test_name,
            "test_method": test_result.test_method,
            "standard": test_result.standard.value,
            "test_date": test_result.test_date.strftime("%Y-%m-%d %H:%M"),
            "operator": test_result.operator,
            "sample_id": test_result.sample_id,
            "manufacturer": test_result.manufacturer,
            "model": test_result.model,
            "serial_number": test_result.serial_number,
            "parameters": test_result.parameters,
            "measurements": test_result.measurements,
            "calculated_values": test_result.calculated_values,
            "pass_fail_criteria": test_result.pass_fail_criteria,
            "overall_result": test_result.overall_result,
            "notes": test_result.notes,
            "interpretation": test_result.notes,  # Will contain AI interpretation
        }

    def _get_default_template(self, report_type) -> str:
        """Get default template ID for report type"""
        template_map = {
            "test_result": "test_result_iec61215",
            "performance": "performance_report",
            "compliance": "compliance_report",
        }
        return template_map.get(report_type.value, "test_result_iec61215")

    async def _generate_output_file(
        self,
        content: str,
        output_format: ReportFormat,
        report_id: str,
        report_title: str,
        output_dir: Path,
        test_results: List[TestResult],
    ) -> Path:
        """
        Generate output file in specified format

        Args:
            content: Report content (markdown)
            output_format: Output format
            report_id: Report ID
            report_title: Report title
            output_dir: Output directory
            test_results: Test results for data export

        Returns:
            Path to generated file
        """
        filename_base = f"{report_id}_{report_title.replace(' ', '_')}"

        if output_format == ReportFormat.PDF:
            return await self._generate_pdf(content, filename_base, output_dir)
        elif output_format == ReportFormat.WORD:
            return await self._generate_word(content, filename_base, output_dir)
        elif output_format == ReportFormat.EXCEL:
            return await self._generate_excel(test_results, filename_base, output_dir)

    async def _generate_pdf(
        self, content: str, filename_base: str, output_dir: Path
    ) -> Path:
        """Generate PDF from markdown content"""
        try:
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration

            # Convert markdown to HTML
            html_content = markdown.markdown(
                content,
                extensions=["tables", "fenced_code", "nl2br"],
            )

            # Wrap in HTML structure
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        line-height: 1.6;
                        margin: 40px;
                        color: #333;
                    }}
                    h1 {{
                        color: #2c3e50;
                        border-bottom: 2px solid #3498db;
                        padding-bottom: 10px;
                    }}
                    h2 {{
                        color: #34495e;
                        margin-top: 30px;
                    }}
                    h3 {{
                        color: #7f8c8d;
                    }}
                    table {{
                        border-collapse: collapse;
                        width: 100%;
                        margin: 20px 0;
                    }}
                    th, td {{
                        border: 1px solid #ddd;
                        padding: 12px;
                        text-align: left;
                    }}
                    th {{
                        background-color: #3498db;
                        color: white;
                    }}
                    strong {{
                        color: #2c3e50;
                    }}
                    hr {{
                        border: none;
                        border-top: 1px solid #ecf0f1;
                        margin: 30px 0;
                    }}
                    @page {{
                        margin: 2.5cm;
                        @bottom-right {{
                            content: "Page " counter(page) " of " counter(pages);
                        }}
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """

            # Generate PDF
            output_path = output_dir / f"{filename_base}.pdf"
            HTML(string=full_html).write_pdf(output_path)

            return output_path

        except Exception as e:
            # Fallback: save as HTML if PDF generation fails
            output_path = output_dir / f"{filename_base}.html"
            with open(output_path, "w") as f:
                f.write(full_html)
            return output_path

    async def _generate_word(
        self, content: str, filename_base: str, output_dir: Path
    ) -> Path:
        """Generate Word document from markdown content"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Parse markdown and add to document
            lines = content.split("\n")
            for line in lines:
                line = line.strip()

                if line.startswith("# "):
                    # Heading 1
                    heading = doc.add_heading(line[2:], level=1)
                    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

                elif line.startswith("## "):
                    # Heading 2
                    doc.add_heading(line[3:], level=2)

                elif line.startswith("### "):
                    # Heading 3
                    doc.add_heading(line[4:], level=3)

                elif line.startswith("---"):
                    # Horizontal rule (page break)
                    doc.add_page_break()

                elif line.startswith("**") and line.endswith("**"):
                    # Bold paragraph
                    p = doc.add_paragraph()
                    p.add_run(line.strip("*")).bold = True

                elif line:
                    # Regular paragraph
                    doc.add_paragraph(line)

            # Save document
            output_path = output_dir / f"{filename_base}.docx"
            doc.save(output_path)

            return output_path

        except Exception as e:
            # Fallback: save as text
            output_path = output_dir / f"{filename_base}.txt"
            with open(output_path, "w") as f:
                f.write(content)
            return output_path

    async def _generate_excel(
        self, test_results: List[TestResult], filename_base: str, output_dir: Path
    ) -> Path:
        """Generate Excel spreadsheet with test data"""
        try:
            import pandas as pd

            # Create Excel writer
            output_path = output_dir / f"{filename_base}.xlsx"

            with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
                # Summary sheet
                summary_data = []
                for test in test_results:
                    summary_data.append({
                        "Test ID": test.test_id,
                        "Test Name": test.test_name,
                        "Sample ID": test.sample_id,
                        "Manufacturer": test.manufacturer,
                        "Model": test.model,
                        "Test Date": test.test_date,
                        "Result": test.overall_result,
                    })

                df_summary = pd.DataFrame(summary_data)
                df_summary.to_excel(writer, sheet_name="Summary", index=False)

                # Individual test sheets
                for i, test in enumerate(test_results[:10]):  # Limit to 10 sheets
                    test_data = {
                        "Parameter": [],
                        "Value": [],
                    }

                    # Add measurements
                    for key, value in test.measurements.items():
                        test_data["Parameter"].append(key)
                        test_data["Value"].append(value)

                    df_test = pd.DataFrame(test_data)
                    sheet_name = f"Test_{i+1}"[:31]  # Excel sheet name limit
                    df_test.to_excel(writer, sheet_name=sheet_name, index=False)

            return output_path

        except Exception as e:
            # Fallback: save as CSV
            output_path = output_dir / f"{filename_base}.csv"
            summary_data = []
            for test in test_results:
                summary_data.append({
                    "Test ID": test.test_id,
                    "Test Name": test.test_name,
                    "Sample ID": test.sample_id,
                    "Result": test.overall_result,
                })

            import pandas as pd
            df = pd.DataFrame(summary_data)
            df.to_csv(output_path, index=False)
            return output_path
